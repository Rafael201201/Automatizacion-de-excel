"""
API Flask mejorada para procesamiento de archivos con opciones de compartir
Incluye integración con múltiples plataformas de compartir y descarga
"""

from flask import Flask, request, jsonify, send_file
from flask_cors import CORS
from werkzeug.utils import secure_filename
import os
import sys
from pathlib import Path
from datetime import datetime
import traceback
import json
import base64
import mimetypes

# Importar librerías de procesamiento
from docx import Document
import pandas as pd
import openpyxl
from openpyxl.styles import Font, PatternFill, Alignment
import re

app = Flask(__name__)
CORS(app, resources={
    r"/api/*": {
        "origins": ["http://localhost:3000", "http://127.0.0.1:3000"],
        "methods": ["GET", "POST", "OPTIONS"],
        "allow_headers": ["Content-Type"],
        "expose_headers": ["Content-Disposition"],
        "supports_credentials": True
    }
})

# Configuración
BASE_DIR = Path(__file__).resolve().parent
UPLOAD_FOLDER = BASE_DIR / 'temp_uploads'
OUTPUT_FOLDER = BASE_DIR / 'temp_outputs'
ALLOWED_EXTENSIONS = {'xlsx', 'xls', 'docx', 'pdf'}

app.config['UPLOAD_FOLDER'] = str(UPLOAD_FOLDER)
app.config['OUTPUT_FOLDER'] = str(OUTPUT_FOLDER)
app.config['MAX_CONTENT_LENGTH'] = 16 * 1024 * 1024  # 16MB

# Crear carpetas
UPLOAD_FOLDER.mkdir(exist_ok=True)
OUTPUT_FOLDER.mkdir(exist_ok=True)

# ==================== UTILIDADES ====================

def allowed_file(filename):
    return '.' in filename and filename.rsplit('.', 1)[1].lower() in ALLOWED_EXTENSIONS

def generate_filename(original_name, action, extension):
    """Genera nombre de archivo único sin la hora"""
    base_name = Path(original_name).stem[:15] # Limita el nombre original a 15 letras
    # Cambiamos esto para que solo tenga Año-Mes-Día
    timestamp = datetime.now().strftime('%Y-%m-%d') 
    return f"{base_name}_{action}_{timestamp}{extension}"

def get_file_metadata(filepath):
    """Obtiene metadata del archivo"""
    path = Path(filepath)
    stats = path.stat()
    
    return {
        'filename': path.name,
        'size': stats.st_size,
        'size_human': format_bytes(stats.st_size),
        'extension': path.suffix,
        'mime_type': mimetypes.guess_type(str(path))[0],
        'created': datetime.fromtimestamp(stats.st_ctime).isoformat(),
        'modified': datetime.fromtimestamp(stats.st_mtime).isoformat()
    }

def format_bytes(bytes):
    """Formatea bytes a formato legible"""
    for unit in ['B', 'KB', 'MB', 'GB']:
        if bytes < 1024.0:
            return f"{bytes:.2f} {unit}"
        bytes /= 1024.0
    return f"{bytes:.2f} TB"

# ==================== PROCESADOR ====================

class FileProcessor:
    """Clase que contiene toda la lógica de procesamiento de archivos"""
    
    @staticmethod
    def get_excel_info(filepath):
        """Obtiene información de un archivo Excel"""
        try:
            df = pd.read_excel(filepath)
            wb = openpyxl.load_workbook(filepath)
            
            return {
                'rows': len(df),
                'columns': len(df.columns),
                'column_names': df.columns.tolist(),
                'sheets': wb.sheetnames,
                'sheet_count': len(wb.sheetnames),
                'has_formulas': any(cell.data_type == 'f' for sheet in wb for row in sheet.rows for cell in row)
            }
        except Exception as e:
            return {'error': str(e)}
    
    @staticmethod
    def get_word_info(filepath):
        """Obtiene información de un archivo Word"""
        try:
            doc = Document(filepath)
            return {
                'paragraphs': len(doc.paragraphs),
                'tables': len(doc.tables),
                'sections': len(doc.sections),
                'total_chars': sum(len(p.text) for p in doc.paragraphs),
                'total_words': sum(len(p.text.split()) for p in doc.paragraphs)
            }
        except Exception as e:
            return {'error': str(e)}
    
    @staticmethod
    def excel_to_word(input_path, output_path, options=None):
        """Convierte Excel a Word con opciones avanzadas"""
        df = pd.read_excel(input_path)
        doc = Document()
        
        # Opciones personalizables
        include_header = options.get('include_header', True) if options else True
        include_stats = options.get('include_stats', True) if options else True
        table_style = options.get('table_style', 'Light Grid Accent 1') if options else 'Light Grid Accent 1'
        
        if include_header:
            doc.add_heading('Datos de Excel', 0)
            doc.add_paragraph(f'Archivo: {Path(input_path).name}')
            doc.add_paragraph(f'Fecha de conversión: {datetime.now().strftime("%Y-%m-%d %H:%M")}')
        
        if include_stats:
            doc.add_heading('Estadísticas', level=2)
            doc.add_paragraph(f'Total de filas: {len(df)}')
            doc.add_paragraph(f'Total de columnas: {len(df.columns)}')
            doc.add_paragraph('')
        
        doc.add_heading('Datos', level=2)
        
        # Crear tabla
        table = doc.add_table(rows=df.shape[0]+1, cols=df.shape[1])
        table.style = table_style
        
        # Encabezados
        for i, col in enumerate(df.columns):
            cell = table.rows[0].cells[i]
            cell.text = str(col)
            cell.paragraphs[0].runs[0].font.bold = True
        
        # Datos
        for i, row in df.iterrows():
            for j, value in enumerate(row):
                table.rows[i+1].cells[j].text = str(value) if pd.notna(value) else ''
        
        doc.save(output_path)
        return output_path
    
    @staticmethod
    def word_to_excel(input_path, output_path, options=None):
        """Convierte Word a Excel con opciones mejoradas"""
        doc = Document(input_path)
        options = options or {}
        
        with pd.ExcelWriter(output_path, engine='openpyxl') as writer:
            # Extraer tablas
            if options.get('extraer_tablas', True) and doc.tables:
                for i, tabla in enumerate(doc.tables):
                    datos = []
                    for fila in tabla.rows:
                        datos.append([celda.text.strip() for celda in fila.cells])
                    
                    if datos:
                        df = pd.DataFrame(datos[1:], columns=datos[0]) if len(datos) > 1 else pd.DataFrame(datos)
                        sheet_name = f'Tabla_{i+1}'[:31]  # Excel limita a 31 caracteres
                        df.to_excel(writer, sheet_name=sheet_name, index=False)
            
            # Extraer texto
            if options.get('extraer_texto', True):
                texto = [p.text for p in doc.paragraphs if p.text.strip()]
                if texto:
                    df_texto = pd.DataFrame({
                        'Línea': range(1, len(texto) + 1),
                        'Contenido': texto
                    })
                    df_texto.to_excel(writer, sheet_name='Texto', index=False)
            
            # Extraer estructura de títulos
            if options.get('extraer_titulos', True):
                titulos = []
                for p in doc.paragraphs:
                    if 'Heading' in p.style.name:
                        titulos.append({
                            'Nivel': p.style.name.replace('Heading ', ''),
                            'Título': p.text
                        })
                if titulos:
                    df_titulos = pd.DataFrame(titulos)
                    df_titulos.to_excel(writer, sheet_name='Estructura', index=False)
            
            # Hoja de metadata
            if options.get('incluir_metadata', False):
                metadata = {
                    'Propiedad': ['Archivo original', 'Fecha conversión', 'Párrafos', 'Tablas', 'Secciones'],
                    'Valor': [
                        Path(input_path).name,
                        datetime.now().strftime('%Y-%m-%d %H:%M'),
                        len(doc.paragraphs),
                        len(doc.tables),
                        len(doc.sections)
                    ]
                }
                df_meta = pd.DataFrame(metadata)
                df_meta.to_excel(writer, sheet_name='Metadata', index=False)
        
        return output_path

processor = FileProcessor()

# ==================== ENDPOINTS ====================

@app.route('/api/health', methods=['GET'])
def health():
    """Verifica que el servidor esté funcionando"""
    return jsonify({
        'status': 'ok',
        'message': 'Servidor Python funcionando',
        'timestamp': datetime.now().isoformat(),
        'version': '2.0'
    })


@app.route('/api/upload', methods=['POST'])
def upload_file():
    """Sube archivo y retorna información preliminar"""
    try:
        if 'file' not in request.files:
            return jsonify({'error': 'No se envió archivo'}), 400
        
        file = request.files['file']
        if file.filename == '':
            return jsonify({'error': 'Nombre de archivo vacío'}), 400
        
        if not allowed_file(file.filename):
            return jsonify({'error': 'Tipo de archivo no permitido'}), 400
        
        filename = secure_filename(file.filename)
        unique_filename = f"{filename}"
        filepath = UPLOAD_FOLDER / unique_filename
        
        file.save(str(filepath))
        
        # Obtener información básica
        file_ext = filename.rsplit('.', 1)[1].lower()
        file_info = get_file_metadata(str(filepath))
        
        # Información específica por tipo
        if file_ext in ['xlsx', 'xls']:
            extra_info = processor.get_excel_info(str(filepath))
            file_info.update(extra_info)
            file_info['type'] = 'excel'
            file_info['available_conversions'] = ['word', 'csv', 'json']
        elif file_ext == 'docx':
            extra_info = processor.get_word_info(str(filepath))
            file_info.update(extra_info)
            file_info['type'] = 'word'
            file_info['available_conversions'] = ['excel']
        
        file_info['unique_filename'] = unique_filename
        file_info['filepath'] = str(filepath)
        
        return jsonify({
            'success': True,
            'message': 'Archivo subido exitosamente',
            'file_info': file_info
        })
    
    except Exception as e:
        return jsonify({
            'error': str(e),
            'traceback': traceback.format_exc()
        }), 500

@app.route('/api/process', methods=['POST'])
def process_file():
    """Procesa el archivo y retorna información para opciones de compartir"""
    try:
        data = request.json
        filepath = data.get('filepath')
        action = data.get('action')
        options = data.get('options', {})
        
        print(f"[PROCESANDO] Archivo: {filepath}")
        print(f"[PROCESANDO] Acción: {action}")
        print(f"[PROCESANDO] Opciones: {options}")
        
        if not filepath or not os.path.exists(filepath):
            print(f"[ERROR] Archivo no encontrado: {filepath}")
            return jsonify({'error': 'Archivo no encontrado'}), 404
        
        input_path = Path(filepath)
        
        # Determinar extensión de salida
        extension_map = {
            'excel_to_word': '.docx',
            'word_to_excel': '.xlsx',
            'excel_to_csv': '.csv',
            'excel_to_json': '.json'
        }
        
        extension = extension_map.get(action, '.xlsx')
        output_filename = generate_filename(input_path.name, action, extension)
        output_path = OUTPUT_FOLDER / output_filename
        
        print(f"[PROCESANDO] Archivo de salida: {output_path}")
        
        # Ejecutar conversión
        try:
            if action == 'excel_to_word':
                processor.excel_to_word(str(input_path), str(output_path), options)
            elif action == 'word_to_excel':
                processor.word_to_excel(str(input_path), str(output_path), options)
            elif action == 'excel_to_csv':
                df = pd.read_excel(str(input_path))
                df.to_csv(str(output_path), index=False, encoding='utf-8-sig')
            elif action == 'excel_to_json':
                df = pd.read_excel(str(input_path))
                df.to_json(str(output_path), orient='records', indent=2, force_ascii=False)
            else:
                print(f"[ERROR] Accion no implementada: {action}")
                return jsonify({'error': f'Acción no implementada: {action}'}), 400
            
            print(f"[EXITO] Archivo procesado: {output_path}")
            
            # Verificar que el archivo se creó
            if not output_path.exists():
                raise Exception("El archivo procesado no se creó correctamente")
            
        except Exception as conv_error:
            print(f"[ERROR] Error en conversion: {str(conv_error)}")
            traceback.print_exc()
            return jsonify({
                'error': f'Error en conversión: {str(conv_error)}',
                'traceback': traceback.format_exc()
            }), 500
        
        # Obtener metadata del archivo procesado
        output_metadata = get_file_metadata(str(output_path))
        
        # Generar URLs de compartir
        share_links = generate_share_links(output_filename, output_metadata)
        
        print(f"[EXITO] Proceso completado exitosamente")
        
        return jsonify({
            'success': True,
            'message': 'Archivo procesado exitosamente',
            'output': {
                'filename': output_filename,
                'path': str(output_path),
                'metadata': output_metadata,
                'download_url': f'/api/download/{output_filename}',
                'share_links': share_links
            }
        })
    
    except Exception as e:
        print(f"[ERROR FATAL] {str(e)}")
        traceback.print_exc()
        return jsonify({
            'error': str(e),
            'traceback': traceback.format_exc()
        }), 500

@app.route('/api/download/<filename>', methods=['GET'])
def download_file(filename):
    """Descarga un archivo procesado"""
    try:
        filepath = OUTPUT_FOLDER / filename
        
        if not filepath.exists():
            return jsonify({'error': 'Archivo no encontrado'}), 404
        
        # Verificar que el archivo está dentro de OUTPUT_FOLDER (seguridad)
        if not str(filepath.resolve()).startswith(str(OUTPUT_FOLDER.resolve())):
            return jsonify({'error': 'Acceso no autorizado'}), 403
        
        return send_file(
            str(filepath),
            as_attachment=True,
            download_name=filename,
            mimetype='application/octet-stream'
        )
    
    except Exception as e:
        print(f"Error en descarga: {str(e)}")
        return jsonify({'error': str(e)}), 500

@app.route('/api/file/<filename>/info', methods=['GET'])
def get_file_info(filename):
    """Obtiene información de un archivo procesado"""
    try:
        filepath = OUTPUT_FOLDER / filename
        
        if not filepath.exists():
            return jsonify({'error': 'Archivo no encontrado'}), 404
        
        metadata = get_file_metadata(str(filepath))
        share_links = generate_share_links(filename, metadata)
        
        return jsonify({
            'success': True,
            'file': metadata,
            'share_links': share_links
        })
    
    except Exception as e:
        return jsonify({'error': str(e)}), 500

@app.route('/api/share/email', methods=['POST'])
def share_via_email():
    """Prepara el archivo para compartir vía email"""
    try:
        data = request.json
        filename = data.get('filename')
        recipient = data.get('recipient')
        
        filepath = OUTPUT_FOLDER / filename
        if not filepath.exists():
            return jsonify({'error': 'Archivo no encontrado'}), 404
        
        # En producción, aquí se enviaría el email
        # Por ahora retornamos un mailto: link
        
        metadata = get_file_metadata(str(filepath))
        subject = f"Archivo compartido: {filename}"
        body = f"Te comparto el archivo {filename} ({metadata['size_human']})"
        
        mailto_link = f"mailto:{recipient}?subject={subject}&body={body}"
        
        return jsonify({
            'success': True,
            'method': 'email',
            'mailto_link': mailto_link,
            'message': 'Abre tu cliente de email para compartir'
        })
    
    except Exception as e:
        return jsonify({'error': str(e)}), 500

@app.route('/api/cleanup', methods=['POST'])
def cleanup():
    """Limpia archivos temporales antiguos"""
    try:
        data = request.json or {}
        max_age_hours = data.get('max_age_hours', 24)
        current_time = datetime.now()
        deleted_count = 0
        
        for folder in [UPLOAD_FOLDER, OUTPUT_FOLDER]:
            for file_path in folder.glob('*'):
                if file_path.is_file():
                    file_time = datetime.fromtimestamp(file_path.stat().st_mtime)
                    age_hours = (current_time - file_time).total_seconds() / 3600
                    
                    if age_hours > max_age_hours:
                        file_path.unlink()
                        deleted_count += 1
        
        return jsonify({
            'success': True,
            'message': f'{deleted_count} archivo(s) eliminado(s)',
            'deleted_count': deleted_count
        })
    
    except Exception as e:
        return jsonify({'error': str(e)}), 500

def generate_share_links(filename, metadata):
    """Genera enlaces para compartir en diferentes plataformas"""
    base_url = "http://localhost:5000"
    download_url = f"{base_url}/api/download/{filename}"
    
    # Texto para compartir
    share_text = f"Archivo: {filename} ({metadata['size_human']})"
    
    return {
        'whatsapp': f"https://wa.me/?text={share_text}%20{download_url}",
        'telegram': f"https://t.me/share/url?url={download_url}&text={share_text}",
        'twitter': f"https://twitter.com/intent/tweet?text={share_text}&url={download_url}",
        'facebook': f"https://www.facebook.com/sharer/sharer.php?u={download_url}",
        'linkedin': f"https://www.linkedin.com/sharing/share-offsite/?url={download_url}",
        'email': f"mailto:?subject={filename}&body={share_text}%20{download_url}",
        'direct': download_url
    }

if __name__ == '__main__':
    print("=" * 70)
    print("Servidor de procesamiento de archivos MEJORADO v2.0")
    print("=" * 70)
    print(f"Carpeta de uploads: {UPLOAD_FOLDER}")
    print(f"Carpeta de outputs: {OUTPUT_FOLDER}")
    print("Servidor corriendo en: http://localhost:5000")
    print("=" * 70)
    print("Caracteristicas:")
    print("   - Conversion Excel <-> Word")
    print("   - Exportacion a CSV/JSON")
    print("   - Compartir en multiples plataformas")
    print("   - Sistema de descarga multiplataforma")
    print("=" * 70)
    
    app.run(debug=True, host='0.0.0.0', port=5000)